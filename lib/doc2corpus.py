"""文档 → 继续预训练（CPT）语料：完整知识无损注入层（零 LLM 成本）。

定位（两段式配方）：
  知识注入层（本模块）：文档原文清洗+分块 → {"text"} 纯文本语料，
    与 minimind/LLaMA-Factory(--stage pt) 等 CPT 流程直接对接；
  表达层（doc2data，后续批次）：基于文档的问答/指令 SFT 数据，CPT 之后使用。

原则：CPT 语料必须保留全部原始知识（清洗只做无损规范化），去重只去重复段落。
"""
from __future__ import annotations

import hashlib
import json
import pathlib
import re
from typing import Any, Dict, Iterator, List

SUPPORTED_EXTS = {".md", ".txt", ".pdf", ".docx"}
PAGE_NUMBER_RE = re.compile(r"^\s*\d{1,4}\s*$")  # 纯页码行（PDF 提取常见噪音）


# ── 导入层 ──────────────────────────────────────────────────────────────────
def import_text(path: str | pathlib.Path) -> str:
    path = pathlib.Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"不支持的文件类型 {ext}（支持 {sorted(SUPPORTED_EXTS)}）")
    if ext in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    raise AssertionError("unreachable")


def clean_text(text: str) -> str:
    """无损规范化：去纯页码行、压缩多余空行、统一换行；不删任何正文内容。"""
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append("")
            continue
        if PAGE_NUMBER_RE.match(stripped):
            continue  # 纯页码行（PDF 提取噪音），非正文
        lines.append(stripped)
    text = "\n".join(lines)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


# ── 分块层 ──────────────────────────────────────────────────────────────────
def chunk_text(text: str, target_chars: int = 2000, overlap: int = 0) -> List[str]:
    """按段落边界分块，目标长度 target_chars；Markdown 标题作为硬边界。
    段落合并直到接近目标长度；overlap>0 时块间回退 N 字符（连续上下文）。
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[str] = []
    buf: List[str] = []
    buf_len = 0
    for para in paragraphs:
        is_heading = para.lstrip().startswith("#")  # Markdown 标题硬边界
        if buf and (is_heading or buf_len + len(para) > target_chars):
            chunks.append("\n\n".join(buf))
            if overlap > 0 and buf:
                tail = buf[-1][-overlap:]
                buf = [tail] if tail else []
                buf_len = len(tail)
            else:
                buf, buf_len = [], 0
        buf.append(para)
        buf_len += len(para)
    if buf:
        chunks.append("\n\n".join(buf))
    return chunks


# ── 去重层 ──────────────────────────────────────────────────────────────────
def chunk_hash(chunk: str) -> str:
    norm = re.sub(r"\s+", "", chunk)
    return hashlib.sha256(norm.encode("utf-8")).hexdigest()[:16]


def doc_to_corpus(
    path: str | pathlib.Path,
    target_chars: int = 2000,
    overlap: int = 0,
    manifest: set[str] | None = None,
) -> Dict[str, Any]:
    """单文档 → 语料条目列表 + 统计。manifest 为跨文档全局查重集合。"""
    manifest = manifest if manifest is not None else set()
    raw = clean_text(import_text(path))
    if not raw:
        return {"entries": [], "stats": {"file": str(path), "chunks": 0, "kept": 0, "dups": 0, "chars": 0}}

    entries: List[Dict[str, str]] = []
    dups = 0
    for i, chunk in enumerate(chunk_text(raw, target_chars, overlap)):
        h = chunk_hash(chunk)
        if h in manifest:
            dups += 1
            continue
        manifest.add(h)
        entries.append({
            "text": chunk,
            "source": pathlib.Path(path).name,
            "chunk_id": f"{pathlib.Path(path).name}#{i}",
        })
    stats = {
        "file": str(path),
        "chunks": len(entries) + dups,
        "kept": len(entries),
        "dups": dups,
        "chars": len(raw),
    }
    return {"entries": entries, "stats": stats}


def write_corpus_jsonl(entries: List[Dict[str, str]], out_path: str | pathlib.Path) -> int:
    """写 CPT 语料（minimind/LLaMA-Factory --stage pt 兼容的 {"text"} 行）。"""
    out_path = pathlib.Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "a", encoding="utf-8") as f:
        for e in entries:
            f.write(json.dumps(e, ensure_ascii=False) + "\n")
    return len(entries)
